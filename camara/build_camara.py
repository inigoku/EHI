#!/usr/bin/env python3
"""Compila una "variación de cámara" (.md) a .docx, .epub y .pdf.

Formato de origen esperado (ver EHI_el_diapason_invisible.md como referencia):
  Línea 1: EL HORIZONTE INTERIOR
  Línea 2: Variación de cámara
  Línea 3: TÍTULO DE LA PIEZA
  Línea 4: Autor
  Líneas 5-6: bajada (dos líneas)
  Luego, separadas por líneas "---": bloques de párrafos separados por
  líneas en blanco. Un párrafo que empieza por "## " es un subtítulo.
  Un párrafo que empieza por uno de los encabezados de sección conocidos
  (OBERTURA, EPÍLOGO, Ensayo:, Lecturas:, Nota del autor, Glosario
  mínimo, Notas y fuentes) es un encabezado de sección. Un párrafo que
  empieza por "*" se renderiza en cursiva completa (notas editoriales).
  Un bloque de varias líneas sin blanco entre ellas conserva los saltos
  de línea (poemas); un bloque de una sola línea es párrafo normal.

Uso:
  python3 build_camara.py EHI_el_diapason_invisible.md
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

SECTION_PREFIXES = (
    "OBERTURA",
    "EPÍLOGO",
    "Ensayo:",
    "Lecturas:",
    "Nota del autor",
    "Glosario mínimo",
    "Notas y fuentes",
)

EMPH_RE = re.compile(r"\*\*([^*]+)\*\*|\*([^*]+)\*")


def is_section_heading(line: str) -> bool:
    return any(line.startswith(p) for p in SECTION_PREFIXES)


def split_emphasis(text: str) -> list[tuple[str, bool, bool]]:
    """Split on **bold** and *italic* markdown into (segment, is_bold, is_italic) triples."""
    parts: list[tuple[str, bool, bool]] = []
    last = 0
    for m in EMPH_RE.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False, False))
        if m.group(1) is not None:
            parts.append((m.group(1), True, False))
        else:
            parts.append((m.group(2), False, True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False, False))
    return parts or [(text, False, False)]


class Block:
    __slots__ = ("kind", "lines")

    def __init__(self, kind: str, lines: list[str]):
        self.kind = kind  # "section" | "subhead" | "italic" | "para"
        self.lines = lines


def parse(path: Path):
    raw = path.read_text(encoding="utf-8")
    header_lines = [l.rstrip() for l in raw.splitlines()[:6]]
    body = "\n".join(raw.splitlines()[7:])  # skip header + blank line

    sections: list[list[Block]] = []
    current: list[Block] = []
    for chunk in re.split(r"\n---\n", body):
        chunk = chunk.strip("\n")
        if not chunk.strip():
            continue
        blocks: list[Block] = []
        for para in re.split(r"\n\s*\n", chunk):
            para = para.strip("\n")
            if not para.strip():
                continue
            lines = [l for l in para.split("\n")]
            first = lines[0].strip()
            if first.startswith("## "):
                blocks.append(Block("subhead", [first[3:].strip()]))
            elif is_section_heading(first):
                blocks.append(Block("section", lines))
            else:
                blocks.append(Block("para", lines))
        sections.append(blocks)
    return header_lines, sections


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def build_docx(header, sections, out_path: Path):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Cm

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Garamond"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def add(text, *, size=11, bold=False, italic=False, align=None, space_before=0, space_after=8, all_caps=False):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text.upper() if all_caps else text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        return p

    def add_multiline(lines, *, italic=False, align=None, size=11):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(8)
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run().add_break()
            for segment, is_bold, is_em in split_emphasis(line):
                if not segment:
                    continue
                r = p.add_run(segment)
                r.italic = italic or is_em
                r.bold = is_bold
                r.font.size = Pt(size)

    # Portada
    add(header[0], size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add(header[1], size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
    add(header[2], size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    add(header[3], size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add(header[4], size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add(header[5], size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    doc.add_page_break()

    for si, blocks in enumerate(sections):
        for block in blocks:
            if block.kind == "section":
                add(block.lines[0], size=15, bold=True, space_before=18, space_after=12)
            elif block.kind == "subhead":
                add(block.lines[0], size=12.5, bold=True, space_before=14, space_after=8)
            else:
                align = None if len(block.lines) > 1 else WD_ALIGN_PARAGRAPH.JUSTIFY
                add_multiline(block.lines, align=align)
        if si < len(sections) - 1:
            doc.add_paragraph().add_run("· · ·").italic = True
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)


# ---------------------------------------------------------------------------
# EPUB
# ---------------------------------------------------------------------------

def build_epub(header, sections, out_path: Path):
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier(f"ehi-camara-{out_path.stem}")
    book.set_title(header[2].title())
    book.set_language("es")
    book.add_author(header[3])

    css = epub.EpubItem(
        uid="style",
        file_name="style/main.css",
        media_type="text/css",
        content=(
            "body{font-family:Georgia,'Liberation Serif',serif;line-height:1.5;"
            "margin:1em 2em;color:#1a1a1a;}"
            "h1{font-size:1.4em;margin-top:1.6em;}"
            "h2{font-size:1.15em;margin-top:1.3em;}"
            "p{margin:0 0 0.9em 0;text-align:justify;}"
            ".poem p{text-align:left;font-style:italic;}"
            ".italic{font-style:italic;}"
            ".cover{text-align:center;}"
            ".cover .kicker{font-style:italic;}"
            ".cover .title{font-size:1.8em;font-weight:bold;margin:0.6em 0;}"
            ".sep{text-align:center;margin:2em 0;}"
        ),
    )
    book.add_item(css)

    def esc(t: str) -> str:
        return (
            t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        )

    def inline_html(text: str) -> str:
        out = []
        for segment, is_bold, is_em in split_emphasis(text):
            escaped = esc(segment)
            if is_bold:
                escaped = f"<strong>{escaped}</strong>"
            if is_em:
                escaped = f"<em>{escaped}</em>"
            out.append(escaped)
        return "".join(out)

    html_parts = [
        '<div class="cover">',
        f"<p class=\"kicker\">{esc(header[0])}<br/>{esc(header[1])}</p>",
        f"<p class=\"title\">{esc(header[2])}</p>",
        f"<p>{esc(header[3])}</p>",
        f"<p class=\"kicker\">{esc(header[4])}<br/>{esc(header[5])}</p>",
        "</div>",
    ]

    for si, blocks in enumerate(sections):
        in_poem = False
        for block in blocks:
            if block.kind == "section":
                html_parts.append(f"<h1>{esc(block.lines[0])}</h1>")
                in_poem = False
            elif block.kind == "subhead":
                html_parts.append(f"<h2>{esc(block.lines[0])}</h2>")
            else:
                html_parts.append(
                    "<p>" + "<br/>".join(inline_html(l) for l in block.lines) + "</p>"
                )
        if si < len(sections) - 1:
            html_parts.append('<p class="sep">· · ·</p>')

    chapter = epub.EpubHtml(title=header[2].title(), file_name="content.xhtml", lang="es")
    chapter.content = "<html><body>" + "\n".join(html_parts) + "</body></html>"
    chapter.add_item(css)
    book.add_item(chapter)

    book.toc = (epub.Link("content.xhtml", header[2].title(), "content"),)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", chapter]

    epub.write_epub(str(out_path), book)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def build_pdf(header, sections, out_path: Path):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

    FONT_DIR = Path("/usr/share/fonts/truetype/liberation")
    pdfmetrics.registerFont(TTFont("Serif", str(FONT_DIR / "LiberationSerif-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("Serif-Bold", str(FONT_DIR / "LiberationSerif-Bold.ttf")))
    pdfmetrics.registerFont(TTFont("Serif-Italic", str(FONT_DIR / "LiberationSerif-Italic.ttf")))

    def esc(t: str) -> str:
        return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    kicker = ParagraphStyle("kicker", fontName="Serif-Italic", fontSize=11, alignment=TA_CENTER, leading=15)
    booktitle = ParagraphStyle("booktitle", fontName="Serif-Bold", fontSize=14, alignment=TA_CENTER, leading=18, spaceAfter=4)
    pdftitle = ParagraphStyle("pdftitle", fontName="Serif-Bold", fontSize=24, alignment=TA_CENTER, leading=30, spaceBefore=14, spaceAfter=14)
    author = ParagraphStyle("author", fontName="Serif", fontSize=13, alignment=TA_CENTER, spaceAfter=22)
    h1 = ParagraphStyle("h1", fontName="Serif-Bold", fontSize=16, spaceBefore=20, spaceAfter=12, leading=20)
    h2 = ParagraphStyle("h2", fontName="Serif-Bold", fontSize=13, spaceBefore=14, spaceAfter=8, leading=16)
    body = ParagraphStyle("body", fontName="Serif", fontSize=11, alignment=TA_JUSTIFY, leading=16, spaceAfter=10)
    poem = ParagraphStyle("poem", fontName="Serif-Italic", fontSize=11, alignment=TA_LEFT, leading=16, spaceAfter=10)
    sep = ParagraphStyle("sep", fontName="Serif-Italic", fontSize=11, alignment=TA_CENTER, spaceBefore=8, spaceAfter=8)

    doc = SimpleDocTemplate(
        str(out_path), pagesize=LETTER,
        leftMargin=2.8 * cm, rightMargin=2.8 * cm, topMargin=2.5 * cm, bottomMargin=2.5 * cm,
    )

    story = []
    story.append(Spacer(1, 3 * cm))
    story.append(Paragraph(f"{esc(header[0])}<br/>{esc(header[1])}", kicker))
    story.append(Spacer(1, 0.6 * cm))
    story.append(Paragraph(esc(header[2]), pdftitle))
    story.append(Paragraph(esc(header[3]), author))
    story.append(Paragraph(f"{esc(header[4])}<br/>{esc(header[5])}", kicker))
    story.append(PageBreak())

    def render(text: str) -> str:
        out = []
        for segment, is_bold, is_em in split_emphasis(text):
            escaped = esc(segment)
            if is_bold:
                escaped = f"<b>{escaped}</b>"
            if is_em:
                escaped = f"<i>{escaped}</i>"
            out.append(escaped)
        return "".join(out)

    for si, blocks in enumerate(sections):
        for block in blocks:
            if block.kind == "section":
                story.append(Paragraph(esc(block.lines[0]), h1))
            elif block.kind == "subhead":
                story.append(Paragraph(esc(block.lines[0]), h2))
            else:
                st = poem if len(block.lines) > 1 else body
                story.append(Paragraph("<br/>".join(render(l) for l in block.lines), st))
        if si < len(sections) - 1:
            story.append(Paragraph("· · ·", sep))

    doc.build(story)


def main():
    if len(sys.argv) != 2:
        print("Uso: build_camara.py archivo.md", file=sys.stderr)
        sys.exit(1)
    src = Path(sys.argv[1])
    header, sections = parse(src)
    stem = src.with_suffix("")
    build_docx(header, sections, stem.with_suffix(".docx"))
    print("docx OK ->", stem.with_suffix(".docx"))
    build_epub(header, sections, stem.with_suffix(".epub"))
    print("epub OK ->", stem.with_suffix(".epub"))
    build_pdf(header, sections, stem.with_suffix(".pdf"))
    print("pdf OK ->", stem.with_suffix(".pdf"))


if __name__ == "__main__":
    main()
