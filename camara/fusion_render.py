"""Renderizadores docx/epub/pdf para las variaciones de cámara y la
Edición de cámara integral.

Trabaja sobre listas de FBlock (build_fusion.FBlock): kind "heading" con
level 1/2/3, o kind "para" con líneas (una sola línea = párrafo normal que
se envuelve; varias líneas = verso, se preservan los saltos). Reutiliza
split_emphasis de build_camara.py para *cursiva* y **negrita** en línea.

Los encabezados de nivel 1 (movimientos / partes) alimentan un índice real
en las tres salidas: estilos de título de Word (navegables en el panel de
Word), tabla de contenidos paginada en el pdf (reportlab TableOfContents),
y navegación multi-entrada en el epub.
Los versos (bloques de varias líneas) admiten sangría espacial: cada línea
de origen puede empezar con espacios en pares (2 espacios = 1 nivel), que
split_verse_line() traduce a indentación tipográfica real en cada formato
en vez de dejarlos como espacios literales.
"""
from __future__ import annotations

from pathlib import Path

from build_camara import split_emphasis


def esc_html(t: str) -> str:
    return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def level1_titles(blocks) -> list[str]:
    return [b.lines[0] for b in blocks if b.kind == "heading" and b.level == 1]


def split_verse_line(line: str) -> tuple[int, str]:
    """('  texto' -> (1, 'texto')); 2 espacios de cabecera = 1 nivel de sangría."""
    stripped = line.lstrip(" ")
    spaces = len(line) - len(stripped)
    return spaces // 2, stripped


def is_poem_title(blocks, i: int) -> bool:
    """Un encabezado de nivel 3 es título de poema si lo sigue un verso
    (bloque de varias líneas): así se distingue de los subtítulos de
    ficción ("I. La primera cerradura"...), que van seguidos de prosa."""
    b = blocks[i]
    if not (b.kind == "heading" and b.level == 3 and i + 1 < len(blocks)):
        return False
    nxt = blocks[i + 1]
    return nxt.kind == "para" and len(nxt.lines) > 1


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def build_docx(header, blocks, out_path: Path):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Garamond"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    INK = RGBColor(0x2A, 0x24, 0x1C)
    RULE = RGBColor(0x9C, 0x8A, 0x6A)

    def add(text, *, size=11, bold=False, italic=False, align=None,
            space_before=0, space_after=8, color=None, style_name=None):
        p = doc.add_paragraph()
        if style_name is not None:
            p.style = doc.styles[style_name]
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
        return p

    def add_rule(p, color=RULE, size_eighths=4, space_pt=6):
        """Añade un filete horizontal bajo el párrafo p (borde inferior)."""
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(size_eighths))
        bottom.set(qn("w:space"), str(space_pt))
        bottom.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_multiline(lines, *, align=None, size=11):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(8)
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run().add_break()
            for segment, is_bold, is_em in split_emphasis(line):
                if not segment:
                    continue
                r = p.add_run(segment)
                r.italic = is_em
                r.bold = is_bold
                r.font.size = Pt(size)

    def add_verse(lines, *, size=11, unit_cm=0.55):
        """Un párrafo por verso, con sangría real según split_verse_line()."""
        n = len(lines)
        for i, raw in enumerate(lines):
            level, text = split_verse_line(raw)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(unit_cm * level)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0 if i < n - 1 else 8)
            for segment, is_bold, is_em in split_emphasis(text):
                if not segment:
                    continue
                r = p.add_run(segment)
                r.italic = is_em
                r.bold = is_bold
                r.font.size = Pt(size)

    HEADING_STYLE = {
        1: dict(size=18, bold=True, space_before=0, space_after=16),
        2: dict(size=13.5, bold=True, space_before=16, space_after=10),
        3: dict(size=11.5, bold=True, space_before=12, space_after=6),
    }
    WORD_STYLE = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}

    # --- Portada ---------------------------------------------------------
    doc.add_paragraph().paragraph_format.space_after = Pt(70)  # empuja hacia el centro vertical
    add(header[0], size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4, color=RULE)
    rule_p = add("", size=1, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_rule(rule_p, space_pt=1)
    add(header[2], size=27, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8, color=INK)
    add(header[1], size=12.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=26)
    add(header[3], size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    if header[4]:
        add(header[4], size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=0, color=RULE)
    if header[5]:
        add(header[5], size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=0, color=RULE)
    doc.add_page_break()

    # --- Índice ------------------------------------------------------------
    titles = level1_titles(blocks)
    if titles:
        add("Índice", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=22)
        for t in titles:
            add(t, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=9, color=INK)
        doc.add_page_break()

    # --- Cuerpo ------------------------------------------------------------
    for i, b in enumerate(blocks):
        if b.kind == "heading":
            if b.level == 1 or is_poem_title(blocks, i):
                doc.add_page_break()
            st = HEADING_STYLE[b.level]
            add(b.lines[0], style_name=WORD_STYLE[b.level], **st)
        elif len(b.lines) > 1:
            add_verse(b.lines)
        else:
            add_multiline(b.lines, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    doc.save(out_path)


# ---------------------------------------------------------------------------
# EPUB
# ---------------------------------------------------------------------------

def build_epub(header, blocks, out_path: Path):
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier("ehi-camara-" + header[2].lower().replace(" ", "-"))
    book.set_title(header[2])
    book.set_language("es")
    book.add_author(header[3])

    css = epub.EpubItem(
        uid="style", file_name="style/main.css", media_type="text/css",
        content=(
            "body{font-family:Georgia,'Liberation Serif',serif;line-height:1.55;"
            "margin:1em 2em;color:#1a1a1a;}"
            "h1{font-size:1.5em;margin-top:2.2em;page-break-before:always;}"
            "h2{font-size:1.2em;margin-top:1.5em;}"
            "h3{font-size:1.05em;margin-top:1.2em;}"
            "h3.poem-title{page-break-before:always;margin-top:2.2em;}"
            "p{margin:0 0 0.9em 0;text-align:justify;}"
            ".cover{text-align:center;padding-top:3em;}"
            ".cover .kicker{font-style:italic;letter-spacing:.08em;"
            "text-transform:uppercase;font-size:.85em;color:#6b5d3f;}"
            ".cover .rule{border:none;border-top:1px solid #9c8a6a;"
            "width:5em;margin:0.9em auto;}"
            ".cover .title{font-size:2em;font-weight:bold;margin:0.3em 0 0.2em;}"
            ".cover .subtitle{font-style:italic;margin:0 0 1.2em;}"
            ".cover .author{font-size:1.05em;margin-top:0.6em;}"
            "nav#toc ol{list-style:none;padding-left:0;}"
            "nav#toc li{margin:0.5em 0;text-align:center;}"
            "nav#toc a{text-decoration:none;color:#1a1a1a;}"
            ".verse{margin:0 0 0.9em 0;font-style:italic;}"
            ".verse p{margin:0;text-align:left;}"
            ".v0{margin-left:0}.v1{margin-left:1.3em}.v2{margin-left:2.6em}"
            ".v3{margin-left:3.9em}.v4{margin-left:5.2em}.v5{margin-left:6.5em}"
            ".v6{margin-left:7.8em}"
        ),
    )
    book.add_item(css)

    def inline_html(text: str) -> str:
        out = []
        for segment, is_bold, is_em in split_emphasis(text):
            escaped = esc_html(segment)
            if is_bold:
                escaped = f"<strong>{escaped}</strong>"
            if is_em:
                escaped = f"<em>{escaped}</em>"
            out.append(escaped)
        return "".join(out)

    html_parts = [
        '<div class="cover">',
        f'<p class="kicker">{esc_html(header[0])}</p>',
        '<hr class="rule"/>',
        f'<p class="title">{esc_html(header[2])}</p>',
        f'<p class="subtitle">{esc_html(header[1])}</p>',
        f'<p class="author">{esc_html(header[3])}</p>',
    ]
    if header[4]:
        html_parts.append(f'<p class="kicker">{esc_html(header[4])}</p>')
    if header[5]:
        html_parts.append(f'<p class="kicker">{esc_html(header[5])}</p>')
    html_parts.append("</div>")

    # Anclas por encabezado de nivel 1, para un índice navegable de verdad.
    toc_links: list = []
    h1_seen = 0
    for i, b in enumerate(blocks):
        if b.kind == "heading":
            if b.level == 1:
                h1_seen += 1
                anchor = f"h1-{h1_seen}"
                html_parts.append(f'<h1 id="{anchor}">{esc_html(b.lines[0])}</h1>')
                toc_links.append(epub.Link(f"content.xhtml#{anchor}", b.lines[0], anchor))
            else:
                tag = f"h{b.level}"
                cls = ' class="poem-title"' if is_poem_title(blocks, i) else ""
                html_parts.append(f"<{tag}{cls}>{esc_html(b.lines[0])}</{tag}>")
        elif len(b.lines) > 1:
            html_parts.append('<div class="verse">')
            for raw in b.lines:
                level, text = split_verse_line(raw)
                html_parts.append(f'<p class="v{min(level, 6)}">{inline_html(text)}</p>')
            html_parts.append("</div>")
        else:
            html_parts.append("<p>" + inline_html(b.lines[0]) + "</p>")

    chapter = epub.EpubHtml(title=header[2], file_name="content.xhtml", lang="es")
    chapter.content = "<html><body>" + "\n".join(html_parts) + "</body></html>"
    chapter.add_item(css)
    book.add_item(chapter)
    book.toc = tuple(toc_links) if toc_links else (epub.Link("content.xhtml", header[2], "content"),)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", chapter]
    epub.write_epub(str(out_path), book)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def build_pdf(header, blocks, out_path: Path):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak, HRFlowable,
    )
    from reportlab.platypus.tableofcontents import TableOfContents

    FONT_DIR = Path("/usr/share/fonts/truetype/liberation")
    try:
        pdfmetrics.registerFont(TTFont("Serif", str(FONT_DIR / "LiberationSerif-Regular.ttf")))
        pdfmetrics.registerFont(TTFont("Serif-Bold", str(FONT_DIR / "LiberationSerif-Bold.ttf")))
        pdfmetrics.registerFont(TTFont("Serif-Italic", str(FONT_DIR / "LiberationSerif-Italic.ttf")))
    except Exception:
        pass  # already registered by a prior call in the same process

    INK = HexColor("#2a241c")
    RULE = HexColor("#9c8a6a")

    def render(text: str) -> str:
        out = []
        for segment, is_bold, is_em in split_emphasis(text):
            escaped = segment.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if is_bold:
                escaped = f"<b>{escaped}</b>"
            if is_em:
                escaped = f"<i>{escaped}</i>"
            out.append(escaped)
        return "".join(out)

    kicker = ParagraphStyle("kicker", fontName="Serif-Italic", fontSize=11, alignment=TA_CENTER,
                             leading=15, textColor=RULE)
    pdftitle = ParagraphStyle("pdftitle", fontName="Serif-Bold", fontSize=25, alignment=TA_CENTER,
                               leading=30, spaceBefore=10, spaceAfter=10, textColor=INK)
    author = ParagraphStyle("author", fontName="Serif", fontSize=13, alignment=TA_CENTER, spaceAfter=22)
    idx_title = ParagraphStyle("idx_title", fontName="Serif-Bold", fontSize=16, alignment=TA_CENTER,
                                spaceBefore=0, spaceAfter=20)
    h1 = ParagraphStyle("h1", fontName="Serif-Bold", fontSize=17, spaceBefore=0, spaceAfter=16, leading=21)
    h2 = ParagraphStyle("h2", fontName="Serif-Bold", fontSize=13.5, spaceBefore=16, spaceAfter=10, leading=17)
    h3 = ParagraphStyle("h3", fontName="Serif-Bold", fontSize=11.5, spaceBefore=12, spaceAfter=6, leading=15)
    body = ParagraphStyle("body", fontName="Serif", fontSize=11, alignment=TA_JUSTIFY, leading=16, spaceAfter=10)
    poem = ParagraphStyle("poem", fontName="Serif-Italic", fontSize=11, alignment=TA_LEFT, leading=16, spaceAfter=10)

    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle("tocH1", fontName="Serif", fontSize=12.5, leading=20,
                        alignment=TA_CENTER, textColor=INK),
    ]
    toc.dotsMinLevel = -1  # sin puntos guía: entradas centradas, más limpio

    class BookDocTemplate(SimpleDocTemplate):
        def afterFlowable(self, flowable):
            if getattr(flowable, "_toc_entry", False):
                text = flowable.getPlainText()
                key = getattr(flowable, "_toc_key")
                self.canv.bookmarkPage(key)
                self.canv.addOutlineEntry(text, key, level=0, closed=False)
                self.notify("TOCEntry", (0, text, self.page, key))

    doc = BookDocTemplate(
        str(out_path), pagesize=LETTER,
        leftMargin=2.8 * cm, rightMargin=2.8 * cm, topMargin=2.5 * cm, bottomMargin=2.5 * cm,
    )

    story = [Spacer(1, 4.5 * cm)]
    story.append(Paragraph(header[0], kicker))
    story.append(HRFlowable(width="15%", thickness=0.75, color=RULE, spaceBefore=10, spaceAfter=14,
                             hAlign="CENTER"))
    story.append(Paragraph(header[2], pdftitle))
    story.append(Paragraph(f"<i>{header[1]}</i>", kicker))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(header[3], author))
    if header[4] or header[5]:
        story.append(Paragraph(f"{header[4]}<br/>{header[5]}", kicker))
    story.append(PageBreak())

    titles = level1_titles(blocks)
    if titles:
        story.append(Paragraph("Índice", idx_title))
        story.append(toc)
        story.append(PageBreak())

    h1_seen = 0
    for bi, b in enumerate(blocks):
        if b.kind == "heading":
            if b.level == 1:
                h1_seen += 1
                story.append(PageBreak())
                p = Paragraph(render(b.lines[0]), h1)
                p._toc_entry = True
                p._toc_key = f"h1-{h1_seen}"
                story.append(p)
            elif b.level == 2:
                story.append(Paragraph(render(b.lines[0]), h2))
            else:
                if is_poem_title(blocks, bi):
                    story.append(PageBreak())
                story.append(Paragraph(render(b.lines[0]), h3))
        elif len(b.lines) > 1:
            n = len(b.lines)
            for i, raw in enumerate(b.lines):
                level, text = split_verse_line(raw)
                line_style = ParagraphStyle(
                    f"verse_{id(b)}_{i}", parent=poem,
                    leftIndent=level * 16, spaceAfter=(0 if i < n - 1 else 10),
                )
                story.append(Paragraph(render(text), line_style))
        else:
            story.append(Paragraph(render(b.lines[0]), body))

    doc.multiBuild(story)
